import docx
import string
from docx.oxml import OxmlElement
from docx.oxml.ns import qn 
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

global_font = 'Arial'
cell_font_size = Pt(12)

def add_text_nl(paragraph, text, bold=False, italic=False, pt=12):
    # Add a run to the paragraph
    run = paragraph.add_run(text + "\n")

    # Add some formatting to the run
    run.bold = bold
    run.italic = italic
    run.font.name = global_font
    run.font.size = docx.shared.Pt(pt)

def add_text(paragraph, text, bold=False, italic=False, pt=12):
    # Add a run to the paragraph
    run = paragraph.add_run(text)

    # Add some formatting to the run
    run.bold = bold
    run.italic = italic
    run.font.name = global_font
    run.font.size = Pt(pt)

# Function to set cell borders
def __set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 0, "val": "single", "color": "#FFFFFF"},
        bottom={"sz": 0, "val": "single", "color": "#FFFFFF"},
        start={"sz": 0, "val": "single", "color": "#FFFFFF"},
        end={"sz": 0, "val": "single", "color": "#FFFFFF"}
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def add_mcq(document, q_idx, question_text, list_of_options):

    # Add another paragraph
    p = document.add_paragraph()
    # Add some formatting to the paragraph
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = 0

    add_text(p, str(q_idx) + ". " + question_text, pt=12)
    table = document.add_table(rows=1, cols=len(list_of_options))

    count = 0
    for chr_op, option in zip(list(map(chr, range(97, 97 + len(list_of_options)))), list_of_options):
        cell = table.cell(0, count)
        # cell.pt = Pt(16)
        cell.text = "(" + chr_op + ") " + option
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        cell.paragraphs[0].line_spacing = 1
        count = count + 1

    # Set the Cell height and width and Hide the table borders
    for row in table.rows:
        for cell in row.cells:
            cell.height = Inches(0.01)
            cell.width = Inches(2)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = cell_font_size
            __set_cell_border(
                cell,
                top={"sz": 0, "val": "single", "color": "#FFFFFF"},
                bottom={"sz": 0, "val": "single", "color": "#FFFFFF"},
                start={"sz": 0, "val": "single", "color": "#FFFFFF"},
                end={"sz": 0, "val": "single", "color": "#FFFFFF"}
            )

# def add_mcq()
# Create a document
doc = docx.Document()

#### Setting the margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(0.5)
    section.right_margin = Cm(0.5)

# Add a paragraph to the document
p = doc.add_paragraph()

# Add some formatting to the paragraph
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.line_spacing = 1.0
p.paragraph_format.space_after = 0

# Add a run to the paragraph
add_text_nl(p, "python-docx", bold=True, italic=True, pt=16)

# Add more text to the same paragraph
add_text_nl(p, "Tutorial", bold=True, pt=16)

# Add another paragraph (left blank for an empty line)
# doc.add_paragraph()


# Add a run and format it


question = "abcdefgh ijklmnopqrstuvwx yzabcdefghijklmnopq rstuvwxyzabcdefghi jklmnopqrstuvwxy zabcdefghijklmn opqrstuvwx yz"
options = ["this is option 1nnnnnnnnnnnnnnnnnnnnnnnnnnnnnnnnnnnnnnn", "this is option 2", "this is option 3", "this is option 4"]
add_mcq(doc, 1, question, options)
add_mcq(doc, 2, question, options)

# Save the document
doc.save("Demo.docx")
